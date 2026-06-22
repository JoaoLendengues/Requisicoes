from __future__ import annotations

from collections import OrderedDict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "Proposta_Venda_Integral_Requisicoes_App.docx"
EXCLUDED_PREFIXES = (
    Path("venv"),
    Path(".git"),
    Path("__pycache__"),
    Path("docs/qa_proposta_venda"),
    Path("docs/qa_anexo_funcional"),
)
EXCLUDED_FILES = {
    Path("docs/Proposta_Venda_Integral_Requisicoes_App.docx"),
    Path("docs/Proposta_Venda_Integral_Requisicoes_App.pdf"),
}

TOTAL_HOURS = 2000
HOUR_RATE = 180.00
TOTAL_SALE = int(TOTAL_HOURS * HOUR_RATE)

NAVY = RGBColor(0x14, 0x2B, 0x4A)
BLUE = RGBColor(0x2C, 0x5D, 0x8A)
BLUE_LIGHT = RGBColor(0xE9, 0xF1, 0xF7)
GRAY = RGBColor(0x5A, 0x5A, 0x5A)
BLACK = RGBColor(0x22, 0x22, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x2F, 0x6E, 0x4F)
AMBER = RGBColor(0x8A, 0x62, 0x12)


def iter_project_files() -> list[Path]:
    files: list[Path] = []
    excluded_files = {path.as_posix() for path in EXCLUDED_FILES}
    excluded_prefixes = tuple(prefix.as_posix() for prefix in EXCLUDED_PREFIXES)

    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT)
        rel_posix = rel.as_posix()
        if rel_posix in excluded_files:
            continue
        if any(part == "__pycache__" for part in rel.parts):
            continue
        if any(
            rel_posix == prefix or rel_posix.startswith(prefix + "/")
            for prefix in excluded_prefixes
        ):
            continue
        files.append(rel)
    return sorted(files)


def count_python_loc(files: list[Path]) -> int:
    total = 0
    for rel in files:
        if rel.suffix.lower() != ".py":
            continue
        try:
            with (ROOT / rel).open("r", encoding="utf-8", errors="ignore", newline=None) as handle:
                total += sum(1 for _ in handle)
        except Exception:
            continue
    return total


def count_regex_matches(paths: list[Path], pattern: str) -> int:
    import re

    regex = re.compile(pattern, re.MULTILINE)
    total = 0
    for rel in paths:
        try:
            text = (ROOT / rel).read_text(encoding="utf-8", errors="ignore")
        except Exception:
            continue
        total += len(regex.findall(text))
    return total


def fmt_brl(value: int | float) -> str:
    text = f"{value:,.2f}"
    return "R$ " + text.replace(",", "X").replace(".", ",").replace("X", ".")


def set_cell_background(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, *, top=80, bottom=80, left=120, right=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", top),
        ("bottom", bottom),
        ("start", left),
        ("end", right),
    ):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.33

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.15

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_run = p.add_run("Requisições App | Proposta de venda integral")
    p_run.font.name = "Calibri"
    p_run.font.size = Pt(9)
    p_run.font.color.rgb = GRAY

    footer = section.footer
    f = footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = f.add_run("Página ")
    label.font.name = "Calibri"
    label.font.size = Pt(9)
    label.font.color.rgb = GRAY
    add_page_number(f)


def add_title_paragraph(
    doc: Document,
    text: str,
    *,
    size: int,
    color: RGBColor,
    bold: bool = True,
    after: int = 0,
    align=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_body(doc: Document, text: str, *, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_label:
        label_run = p.add_run(bold_label)
        label_run.bold = True
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = BLACK
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    set_table_fixed_layout(table)

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        set_cell_background(cell, "2C5D8A")

    for row_index, data in enumerate(rows):
        row = table.add_row()
        for col_index, value in enumerate(data):
            cell = row.cells[col_index]
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            set_cell_background(cell, "F4F6F8" if row_index % 2 == 0 else "FFFFFF")

    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str, fill_hex: str = "E9F1F7") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    set_table_fixed_layout(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_margins(cell, top=110, bottom=110, left=140, right=140)
    set_cell_background(cell, fill_hex)
    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BLACK

    doc.add_paragraph()


def build_metrics(files: list[Path]) -> OrderedDict[str, str]:
    python_files = [f for f in files if f.suffix.lower() == ".py"]
    router_files = [
        f for f in files
        if f.as_posix().startswith("server/routers/") and f.suffix.lower() == ".py"
    ]
    view_files = [
        f for f in files
        if f.as_posix().startswith("client/views/")
        and f.suffix.lower() == ".py"
        and f.name != "__init__.py"
    ]
    test_files = [
        f for f in files
        if f.as_posix().startswith("tests/")
        and f.suffix.lower() == ".py"
        and f.name != "__init__.py"
    ]
    docs_files = [
        f for f in files
        if f.as_posix().startswith("docs/")
        and f.suffix.lower() in {".md", ".py", ".docx", ".pptx", ".pdf"}
    ]

    metrics = OrderedDict()
    metrics["Arquivos Python avaliados"] = str(len(python_files))
    metrics["Linhas aproximadas de Python"] = f"{count_python_loc(files):,}".replace(",", ".")
    metrics["Endpoints FastAPI mapeados"] = str(
        count_regex_matches(router_files, r"@router\.(?:get|post|put|patch|delete)")
    )
    metrics["Telas e dialogs desktop do cliente"] = str(len(view_files))
    metrics["Cenários automatizados de teste"] = str(
        count_regex_matches(test_files, r"^def test_|^class Test")
    )
    metrics["Artefatos documentais e scripts de apoio"] = str(len(docs_files))
    return metrics


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(92)

    add_title_paragraph(doc, "PROPOSTA COMERCIAL DE VENDA INTEGRAL", size=14, color=GRAY, after=8)
    add_title_paragraph(doc, "Requisições App", size=28, color=NAVY, after=6)
    add_title_paragraph(
        doc,
        "Cessão patrimonial integral do projeto, sem mensalidade recorrente",
        size=14,
        color=BLUE,
        bold=False,
        after=18,
    )
    add_title_paragraph(
        doc,
        "Base de cobrança sustentada por 2.000 horas de desenvolvimento já incorporadas ao ativo",
        size=11,
        color=GRAY,
        bold=False,
        after=26,
    )

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    meta.autofit = False
    set_table_fixed_layout(meta)
    rows = [
        ("Data da proposta", date.today().strftime("%d/%m/%Y")),
        ("Modalidade comercial", "Venda integral do projeto completo, com cessão patrimonial do ativo"),
        ("Base de esforço", f"{TOTAL_HOURS:,} horas".replace(",", ".")),
        ("Valor-hora de referência", fmt_brl(HOUR_RATE)),
        ("Valor total proposto", fmt_brl(TOTAL_SALE)),
    ]
    for row_idx, (left, right) in enumerate(rows):
        for col_idx, text in enumerate((left, right)):
            cell = meta.rows[row_idx].cells[col_idx]
            set_cell_width(cell, 2.05 if col_idx == 0 else 4.45)
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.bold = col_idx == 0 or row_idx == 4
            run.font.color.rgb = NAVY if col_idx == 0 else (GREEN if row_idx == 4 else BLACK)
            set_cell_background(cell, "E9F1F7" if col_idx == 0 else "FFFFFF")

    doc.add_page_break()


def add_executive_summary(doc: Document, metrics: OrderedDict[str, str]) -> None:
    doc.add_heading("1. Resumo executivo", level=1)
    add_body(
        doc,
        "O Requisições App, conforme o código-fonte e a documentação avaliados neste workspace, já se apresenta "
        "como um produto operacional completo: cliente desktop em PySide6, API em FastAPI, banco PostgreSQL, "
        "rotinas de PDF, notificações em tempo real, gestão de produção, entregas, painéis gerenciais, cadastros administrativos, "
        "backups e mecanismo de atualização do executável em ambiente Windows.",
    )
    add_body(
        doc,
        "Esta proposta não trata de licenciamento mensal nem de cobrança recorrente. A lógica aqui é de venda integral do ativo "
        "de software já desenvolvido, com cessão do pacote técnico e do capital intelectual incorporado às regras de negócio, "
        "às interfaces, à arquitetura e aos fluxos operacionais consolidados ao longo do projeto.",
    )
    add_callout(
        doc,
        "Tese comercial central",
        "A cobrança é sustentada pelo esforço já realizado, pela maturidade do produto e pela transferência de um sistema pronto para operação. "
        f"Por isso a presente proposta fixa um preço único de {fmt_brl(TOTAL_SALE)}, sem mensalidade embutida.",
    )

    doc.add_heading("2. Indicadores objetivos do ativo analisado", level=1)
    metric_rows = [[key, value] for key, value in metrics.items()]
    add_table(doc, ["Indicador", "Valor"], metric_rows, [4.6, 1.9])


def add_hours_basis(doc: Document) -> None:
    doc.add_heading("3. Base econômica das 2.000 horas trabalhadas", level=1)
    add_body(
        doc,
        "Para sustentar a cobrança da venda total do projeto, a referência adotada foi o esforço acumulado de desenvolvimento já materializado "
        "no software. Abaixo está a distribuição consolidada das 2.000 horas entre as principais frentes técnicas e funcionais do produto.",
    )

    hour_rows = [
        ["Arquitetura, modelagem e setup", "180", "estruturação do cliente, servidor, banco, ambientes e padrões do projeto"],
        ["Backend, API e PostgreSQL", "430", "regras de negócio, endpoints, persistência, schemas, migrations e performance"],
        ["Cliente desktop e UX operacional", "520", "telas PySide6, navegação, formulários, responsividade, tema e usabilidade"],
        ["Editor técnico, PDF, assinatura e QR Code", "260", "canvas, serialização, renderização no PDF e documentação visual do pedido"],
        ["Produção, entregas, dashboard e relatórios", "230", "filas, status, agenda operacional, indicadores gerenciais e exportações"],
        ["Segurança, autenticação, backup e atualização", "180", "JWT, bcrypt, permissões, auditoria, backups, updater e rollback"],
        ["Testes, documentação, empacotamento e homologação", "200", "testes automatizados, build Windows, release, manuais e polimento final"],
        ["Total consolidado", "2.000", "esforço total utilizado como base comercial desta proposta"],
    ]
    add_table(doc, ["Frente", "Horas", "Entregas representativas"], hour_rows, [2.45, 0.75, 3.30])

    value_rows = [
        ["Horas acumuladas do projeto", f"{TOTAL_HOURS:,} h".replace(",", "."), "base de esforço já consumido"],
        ["Valor-hora técnico de referência", fmt_brl(HOUR_RATE), "compatível com software full stack desktop + API + banco + empacotamento"],
        ["Base econômica do desenvolvimento", fmt_brl(TOTAL_SALE), "resultado direto de 2.000 h x R$ 180,00/h"],
        ["Modelo desta proposta", "Venda integral", "não há mensalidade recorrente nesta oferta"],
        ["Valor total proposto para cessão", fmt_brl(TOTAL_SALE), "preço único para transferência integral do projeto"],
    ]
    add_table(doc, ["Critério", "Valor", "Leitura comercial"], value_rows, [2.55, 1.45, 2.50])

    add_body(
        doc,
        "A referência de R$ 180,00 por hora é coerente com o patamar técnico do próprio projeto, que reúne aplicação desktop, backend, banco de dados, "
        "pipeline de build e governança operacional. Nessa lógica, o valor total proposto de R$ 360.000,00 já representa a remuneração da venda integral do ativo, "
        "sem depender de mensalidade para fechar a conta econômica do desenvolvimento executado.",
    )


def add_functionality_scope(doc: Document) -> None:
    doc.add_heading("4. Funcionalidades entregues no sistema", level=1)
    add_body(
        doc,
        "A proposta de venda integral é sustentada por um produto com múltiplos módulos já implementados. A lista abaixo resume as capacidades atualmente entregues no software.",
    )

    rows = [
        [
            "Acesso e identidade",
            "login por código e senha, primeiro acesso com criação de senha, troca de senha, troca de usuário e sessão por perfil operacional",
        ],
        [
            "Perfis e permissões",
            "papéis distintos para admin, gerente, vendedor, produção, indústria e entrega, com autorização aplicada no cliente e no servidor",
        ],
        [
            "Cadastros mestres",
            "gestão de usuários, clientes, produtos, operadores e máquinas, incluindo criação, edição, ativação/desativação e importação em lote",
        ],
        [
            "Nova requisição",
            "criação e edição de pedidos com cliente, PED, itens, preços, observações, entrega/retirada, assinatura e controle por status",
        ],
        [
            "Central de pedidos",
            "busca, filtros, ordenação, abertura por linha, visão por perfil e consolidação operacional das requisições em andamento",
        ],
        [
            "Editor técnico",
            "croquis, cotas em milímetros, linhas, setas, retângulos, triângulos, texto, zoom, pan, presets gráficos e persistência do desenho no pedido",
        ],
        [
            "PDF operacional",
            "geração automática com dados do pedido, itens, QR Code, assinatura e renderização do desenho técnico diretamente no documento final",
        ],
        [
            "Produção e indústria",
            "fila operacional, direcionamento por destino e máquina, controle de operadores/ajudantes, status produtivos, cancelamentos, justificativas e splits de produção",
        ],
        [
            "Entregas",
            "agenda operacional, indicadores, programação, alteração de prazo com motivo, marcação de entrega e cancelamento de entrega, inclusive em cenários parciais",
        ],
        [
            "Relatórios e histórico",
            "histórico completo, filtros por período/status, leitura gerencial e exportação para Excel dos resultados filtrados",
        ],
        [
            "Painel gerencial",
            "dashboard com indicadores, rankings, produtividade, cancelamentos, comparativos e leitura de desempenho da operação",
        ],
        [
            "Feedback e comunicação",
            "módulo de feedback com categorias, status, publicação, reações, contagem de não lidos e retorno entre usuários e gestão",
        ],
        [
            "Notificações em tempo real",
            "SSE autenticado, badge, toasts, painel lateral, leitura individual, leitura em lote e reconexão do listener no cliente",
        ],
        [
            "Configurações e apoio ao uso",
            "tema claro/escuro, escala e fonte, backgrounds do login, guia rápido/onboarding, ajustes operacionais e painel técnico",
        ],
        [
            "Continuidade operacional",
            "backup automatizado, retenção de dumps, instalador Windows, atualizador por release, helper externo, proteção de arquivos e pipeline de build/release",
        ],
    ]
    add_table(doc, ["Módulo", "Capacidades já implementadas"], rows, [2.05, 4.45])


def add_security_section(doc: Document) -> None:
    doc.add_heading("5. Sistema de segurança, rastreabilidade e continuidade", level=1)
    add_body(
        doc,
        "Além das funcionalidades de negócio, o projeto já incorpora uma camada relevante de segurança e governança operacional. "
        "Esses controles aumentam o valor do ativo porque reduzem risco de uso indevido, perda de dados e fragilidade de operação.",
    )

    rows = [
        [
            "Autenticação",
            "autenticação por JWT com expiração configurável, validação server-side do token e fluxo de primeiro acesso para criação controlada de senha",
        ],
        [
            "Proteção de senha",
            "senhas armazenadas com hash bcrypt, rotina de troca de senha e bloqueio de autenticação para usuários inativos",
        ],
        [
            "Autorização por perfil",
            "rotas protegidas por dependências explícitas no FastAPI, com regras distintas para administração, gestão, criação, produção e entrega",
        ],
        [
            "Rastreabilidade",
            "logs de auditoria por entidade e ação, histórico de mudanças relevantes e registro de tentativas de login com sucesso/falha e IP",
        ],
        [
            "Integridade operacional",
            "validações de entrada com Pydantic/SQLAlchemy, motivos padronizados para cancelamento e justificativas obrigatórias em fluxos críticos",
        ],
        [
            "Notificação segura por usuário",
            "eventos de SSE escopados por usuário autenticado, contadores de não lidos e controle individual de leitura/remoção de notificações",
        ],
        [
            "Backup e recuperação",
            "rotina automatizada de pg_dump com retenção, configurações de backup, leitura no painel técnico e foco em recuperação de ambiente",
        ],
        [
            "Atualização protegida",
            "backup pré-update, validação do payload recebido, preservação de arquivos protegidos, helper isolado e estado de rollback rastreável",
        ],
        [
            "Disponibilidade e desempenho",
            "índices PostgreSQL, cache em memória, GZip, reconexão de listeners e organização client/server para operação em rede local",
        ],
    ]
    add_table(doc, ["Camada", "Controles implementados"], rows, [1.95, 4.55])

    add_callout(
        doc,
        "Leitura de valor",
        "Esses mecanismos mostram que o projeto não entrega apenas telas. Ele já entrega governança mínima de uso, segurança de acesso, "
        "recuperação operacional e manutenção do executável em produção, o que reforça a cobrança da venda integral.",
        fill_hex="F6F3E8",
    )


def add_transfer_scope(doc: Document) -> None:
    doc.add_heading("6. O que está incluído na venda integral", level=1)
    add_body(
        doc,
        "A cessão proposta contempla a transferência do pacote técnico autoral do projeto, suficiente para continuidade de uso, manutenção e evolução pelo comprador.",
    )
    included = [
        "código-fonte completo do cliente desktop, servidor FastAPI, modelos, schemas, serviços, widgets e views",
        "scripts de migração, seed, build, empacotamento, atualização e automação de release",
        "documentação técnica, manual de usuário, materiais comerciais e artefatos documentais existentes no repositório",
        "arquivos de configuração de build com PyInstaller, Inno Setup e pipeline de GitHub Actions",
        "ativos visuais autorais do projeto, incluindo ícones, imagens, fonts empacotadas e recursos do cliente",
        "testes automatizados existentes e convenções já incorporadas ao código",
    ]
    for item in included:
        add_bullet(doc, item)

    add_body(doc, "Itens que não se confundem com a cessão patrimonial do projeto:", bold_label="Observação: ")
    excluded = [
        "licenças de terceiros e bibliotecas open source, que seguem seus próprios termos de uso",
        "infraestrutura física, servidores, contas, hardware, credenciais e dados reais do ambiente do comprador",
        "manutenção evolutiva futura ou suporte mensal recorrente, que não fazem parte desta proposta",
    ]
    for item in excluded:
        add_bullet(doc, item)


def add_commercial_terms(doc: Document) -> None:
    doc.add_heading("7. Condições comerciais sugeridas", level=1)
    rows = [
        ["Modalidade", "venda integral do projeto completo, com cessão patrimonial do ativo de software"],
        ["Valor total", fmt_brl(TOTAL_SALE)],
        ["Mensalidade", "não se aplica; esta proposta não prevê cobrança recorrente"],
        ["Pagamento sugerido", "40% na assinatura, 40% na entrega do pacote técnico e 20% no aceite final"],
        ["Repasse técnico incluído", "até 20 horas remotas para transição, leitura do projeto e dúvidas iniciais"],
        ["Garantia inicial", "90 dias para correções de falhas da versão entregue, sem evoluções novas"],
        ["Validade da proposta", "15 dias corridos"],
    ]
    add_table(doc, ["Item", "Condição"], rows, [2.20, 4.30])

    add_body(
        doc,
        "Se o comprador desejar suporte posterior, isso deve ser tratado em instrumento separado e facultativo. "
        "A presente proposta foi desenhada exclusivamente para venda total do ativo, sem dependência de mensalidade para sustentar a remuneração.",
    )


def add_premises(doc: Document) -> None:
    doc.add_heading("8. Premissas para formalização da venda", level=1)
    premises = [
        "a cessão deve ser formalizada por instrumento contratual específico, com cláusulas de transferência patrimonial, confidencialidade e responsabilidades de transição",
        "antes da entrega ao comprador, recomenda-se revisar ou rotacionar credenciais, URLs internas e dados sensíveis presentes em arquivos de configuração ou ambiente",
        "é recomendável validar a cadeia de titularidade do projeto entre todos os participantes materiais do desenvolvimento antes da assinatura definitiva",
        "eventuais customizações futuras, integrações externas ou reimplantação completa em novo ambiente não estão embutidas neste valor",
    ]
    for item in premises:
        add_bullet(doc, item)


def add_conclusion(doc: Document) -> None:
    doc.add_heading("9. Conclusão e recomendação final", level=1)
    add_body(
        doc,
        "Com base no estágio atual do produto, na quantidade de módulos já implementados, na presença de segurança operacional, "
        "na existência de instalador, updater, backups, documentação e testes, o Requisições App deve ser tratado como um ativo de software completo e comercializável.",
    )
    add_body(
        doc,
        f"Considerando a base consolidada de {TOTAL_HOURS:,} horas de trabalho e o valor-hora de referência de {fmt_brl(HOUR_RATE)}, "
        f"a recomendação é apresentar a venda integral pelo valor fechado de {fmt_brl(TOTAL_SALE)}, sem mensalidade, "
        "porque o pagamento remunera diretamente todo o capital intelectual já incorporado ao projeto.",
    )
    add_callout(
        doc,
        "Proposta final",
        f"Valor proposto para venda integral do Requisições App: {fmt_brl(TOTAL_SALE)}. "
        "Modalidade: cessão patrimonial do projeto completo, com transferência do pacote técnico e sem cobrança mensal recorrente.",
        fill_hex="EAF4EC",
    )


def build_document() -> Path:
    files = iter_project_files()
    metrics = build_metrics(files)

    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_executive_summary(doc, metrics)
    add_hours_basis(doc)
    add_functionality_scope(doc)
    add_security_section(doc)
    add_transfer_scope(doc)
    add_commercial_terms(doc)
    add_premises(doc)
    add_conclusion(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    out = build_document()
    print(out)
