from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_sale_proposal import (
    BLACK,
    BLUE,
    GRAY,
    NAVY,
    ROOT,
    add_body,
    add_bullet,
    add_callout,
    add_number,
    add_table,
    add_title_paragraph,
    configure_styles,
)


OUT_DOCX = ROOT / "docs" / "Anexo_Comercial_Funcional_Requisicoes_App.docx"


def _set_header_text(doc: Document, text: str) -> None:
    section = doc.sections[0]
    paragraph = section.header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(95)

    add_title_paragraph(doc, "ANEXO COMERCIAL FUNCIONAL", size=14, color=GRAY, after=8)
    add_title_paragraph(doc, "Requisições App", size=28, color=NAVY, after=6)
    add_title_paragraph(
        doc,
        "Visão não técnica do sistema: telas, processos e controles operacionais",
        size=13,
        color=BLUE,
        bold=False,
        after=22,
    )
    add_title_paragraph(
        doc,
        "Material de apoio para apresentação comercial e entendimento do funcionamento do produto",
        size=11,
        color=GRAY,
        bold=False,
        after=26,
    )
    doc.add_page_break()


def add_overview(doc: Document) -> None:
    doc.add_heading("1. O que é o sistema", level=1)
    add_body(
        doc,
        "O Requisições App é um sistema interno de gestão operacional que organiza o fluxo de pedidos "
        "desde a criação da requisição até o acompanhamento de produção, entrega, faturamento e histórico."
    )
    add_body(
        doc,
        "Na prática, ele conecta as áreas comercial, gestão, produção, indústria, entrega e administração "
        "em um único ambiente de trabalho, substituindo controles espalhados em papel, conversa informal e planilhas soltas."
    )
    add_callout(
        doc,
        "Leitura comercial simples",
        "É um sistema para registrar pedidos com clareza, acompanhar o que está acontecendo com cada requisição "
        "e reduzir retrabalho, perda de informação e dependência de memória das equipes.",
    )


def add_audience(doc: Document) -> None:
    doc.add_heading("2. Quem usa o sistema", level=1)
    audiences = [
        ["Administração", "configura regras, usuários, cadastros, parâmetros e visão geral do ambiente"],
        ["Gestão", "acompanha indicadores, atrasos, produtividade, histórico e situação geral da operação"],
        ["Vendas", "cria requisições, acompanha andamento, ajusta informações e consulta histórico"],
        ["Produção / A&R", "recebe pedidos, organiza fila, direciona produção e atualiza o andamento"],
        ["Indústria", "acompanha e executa etapas produtivas vinculadas às requisições"],
        ["Entrega", "visualiza o que precisa sair, controla entregas e confirma o que já foi concluído"],
    ]
    add_table(doc, ["Perfil", "Papel no dia a dia"], audiences, [1.8, 4.7])


def add_screens(doc: Document) -> None:
    doc.add_heading("3. Telas e módulos do sistema", level=1)
    add_body(doc, "A solução é organizada em módulos de uso diário e módulos de apoio administrativo.")

    doc.add_heading("3.1 Telas operacionais", level=2)
    operational_rows = [
        ["Login e primeiro acesso", "entrada no sistema, definição de senha inicial e troca de usuário", "todos"],
        ["Tela principal", "menu lateral com acesso aos módulos disponíveis conforme o perfil", "todos"],
        ["Dashboard gerencial", "visão consolidada de pedidos, produção, atrasos, cancelamentos e alertas", "gestão/admin"],
        ["Central de pedidos", "consulta geral das requisições com busca, filtros, ordenação e abertura de pedido", "todos"],
        ["Formulário da requisição", "criação e edição de pedidos com cliente, itens, prazo, observações e entrega/retirada", "vendas/gestão/admin"],
        ["Editor de desenho", "criação de croquis, medidas e anotações visuais para detalhar o pedido", "vendas/gestão/admin"],
        ["Produção / A&R / Indústria", "recebimento dos pedidos, fila operacional, acompanhamento por máquina e andamento produtivo", "produção/indústria/gestão/admin"],
        ["Central de entregas", "controle do que precisa ser entregue e confirmação do que já saiu", "entrega/gestão/admin"],
        ["Histórico", "consulta completa das requisições já registradas no sistema", "todos"],
        ["Feedback", "registro de sugestões, melhorias e problemas percebidos pelos usuários", "todos"],
    ]
    add_table(doc, ["Tela", "O que faz", "Uso principal"], operational_rows, [1.65, 3.55, 1.3])

    doc.add_heading("3.2 Telas administrativas e de apoio", level=2)
    support_rows = [
        ["Configurações", "preferências do usuário, senha, aparência, ajuda e parâmetros operacionais"],
        ["Gestão de usuários", "cadastro, edição, permissão e organização dos acessos"],
        ["Gestão de clientes", "cadastro, busca, manutenção e importação de clientes"],
        ["Gestão de máquinas", "cadastro das máquinas e acompanhamento de disponibilidade"],
        ["Gestão de operadores", "cadastro da equipe operacional usada no fluxo produtivo"],
        ["Atualizações do sistema", "consulta e aplicação de novas versões do aplicativo"],
        ["Guia rápido", "tour orientado para ensinar o uso das telas e acelerar a adoção"],
        ["Painel de sistema", "visão administrativa de funcionamento, controles e suporte interno"],
    ]
    add_table(doc, ["Tela", "Finalidade"], support_rows, [2.0, 4.5])


def add_processes(doc: Document) -> None:
    doc.add_heading("4. Processos atendidos pelo sistema", level=1)

    doc.add_heading("4.1 Fluxo comercial do pedido", level=2)
    commercial_steps = [
        "o usuário faz login no sistema conforme seu perfil",
        "abre a central de pedidos e inicia uma nova requisição",
        "localiza o cliente por nome, código ou documento",
        "informa número do pedido, forma de atendimento e prazo",
        "adiciona os itens, quantidades, observações e demais detalhes",
        "inclui, se necessário, um desenho ou croqui de apoio",
        "salva a requisição e gera o documento final do pedido",
        "acompanha o status do pedido até a conclusão",
    ]
    for item in commercial_steps:
        add_number(doc, item)

    doc.add_heading("4.2 Fluxo de produção", level=2)
    production_steps = [
        "a produção recebe a requisição encaminhada pelo comercial",
        "o pedido entra em fila e pode ser distribuído por destino ou máquina",
        "a equipe acompanha o andamento e atualiza o status do pedido",
        "casos de mudança de prazo podem ser devolvidos com justificativa",
        "quando a etapa produtiva termina, o sistema registra a evolução para as áreas seguintes",
    ]
    for item in production_steps:
        add_number(doc, item)

    doc.add_heading("4.3 Fluxo de entrega e encerramento", level=2)
    delivery_steps = [
        "a área de entrega visualiza o que está pronto ou programado para sair",
        "o pedido pode ser marcado como entregue conforme a execução",
        "a gestão e o comercial passam a enxergar a situação atual em tempo real",
        "o pedido finalizado permanece disponível no histórico para consulta futura",
    ]
    for item in delivery_steps:
        add_number(doc, item)

    doc.add_heading("4.4 Processos administrativos", level=2)
    admin_processes = [
        "cadastro e manutenção de usuários",
        "cadastro e importação de clientes, produtos e operadores",
        "controle de máquinas e parâmetros de operação",
        "coleta de feedback das equipes",
        "manutenção de regras de prazo, cancelamento e rotinas operacionais",
    ]
    for item in admin_processes:
        add_bullet(doc, item)


def add_controls(doc: Document) -> None:
    doc.add_heading("5. Controles de negócio e governança operacional", level=1)
    controls = [
        "perfis de acesso diferentes por função, limitando o que cada usuário pode visualizar ou alterar",
        "rastreabilidade de cada requisição ao longo do fluxo operacional",
        "registro de mudanças de status e justificativas quando o prazo precisa ser alterado",
        "cancelamento com motivo padronizado, reduzindo perda de contexto",
        "histórico consolidado para auditoria interna e consulta futura",
        "notificações automáticas sobre eventos importantes do processo",
        "alertas para pedidos próximos do prazo ou em atraso",
        "exportação de informações para análise e acompanhamento gerencial",
        "ajuda guiada para onboarding de novos usuários e padronização do uso",
    ]
    for item in controls:
        add_bullet(doc, item)


def add_information_bases(doc: Document) -> None:
    doc.add_heading("6. Bases de informação mantidas no sistema", level=1)
    bases = [
        ["Usuários", "quem acessa, qual perfil possui e quais permissões operacionais carrega"],
        ["Clientes", "base comercial usada para abertura das requisições"],
        ["Produtos", "itens usados na montagem dos pedidos"],
        ["Operadores", "equipe operacional usada na produção"],
        ["Máquinas", "recursos produtivos usados para organização e acompanhamento"],
        ["Requisições", "pedidos, prazos, observações, status e histórico associado"],
        ["Motivos operacionais", "regras de cancelamento, justificativas e parâmetros do fluxo"],
        ["Feedbacks", "sugestões e ocorrências enviadas pelos usuários do sistema"],
    ]
    add_table(doc, ["Base", "Uso prático"], bases, [1.6, 4.9])


def add_benefits(doc: Document) -> None:
    doc.add_heading("7. Valor operacional entregue pelo sistema", level=1)
    benefits = [
        "centraliza o processo em um único ambiente de trabalho",
        "reduz ruído entre comercial, produção e entrega",
        "diminui retrabalho por falta de informação ou pedido incompleto",
        "melhora previsibilidade e acompanhamento de prazo",
        "permite histórico pesquisável de tudo o que foi solicitado",
        "facilita cobrança gerencial e leitura rápida da operação",
        "acelera treinamento de novos usuários por já ter fluxo orientado",
        "reduz dependência de papel, planilha paralela e memória individual",
    ]
    for item in benefits:
        add_bullet(doc, item)

    add_callout(
        doc,
        "Resumo final",
        "Sem entrar em tecnologia, o sistema já cobre cadastro, operação, acompanhamento, controle, "
        "gestão e histórico do processo de requisições. Isso o torna um produto pronto para uso real, "
        "e não apenas um conceito ou uma tela isolada.",
    )


def add_closing(doc: Document) -> None:
    doc.add_heading("8. Fechamento comercial", level=1)
    add_body(
        doc,
        "Este anexo foi estruturado para apresentar o sistema pela ótica funcional e operacional. "
        "Ele pode ser usado como material de apoio em proposta comercial, apresentação executiva ou negociação de venda do ativo."
    )


def build_document() -> Path:
    doc = Document()
    configure_styles(doc)
    _set_header_text(doc, "Requisições App | Anexo comercial funcional")
    add_cover(doc)
    add_overview(doc)
    add_audience(doc)
    add_screens(doc)
    add_processes(doc)
    add_controls(doc)
    add_information_bases(doc)
    add_benefits(doc)
    add_closing(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    out = build_document()
    print(out)
